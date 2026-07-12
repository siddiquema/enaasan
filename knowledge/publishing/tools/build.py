#!/usr/bin/env python3
"""
Enaasan Knowledge Layer — single-source builder.

Reads a publication's Markdown source (with YAML front matter, following
knowledge/publishing/markdown-standard.md) and renders:

  output/html/<document-id>.html   — standalone, screen + print CSS
  output/pdf/<document-id>.pdf     — via Microsoft Edge headless (A4, print-friendly)
  output/docx/<document-id>.docx   — via python-docx, Google-Docs-import-safe

Usage:
  python build.py <path-to-source.md> [--skip-pdf]

Only the Markdown constructs in markdown-standard.md are supported:
headings (#/##/###), paragraphs, **bold**, *italic*, bullet lists,
hand-numbered lists, task lists (- [ ]), pipe tables, blockquote callouts,
fill-in lines (underscores), and --- rules.
"""
import io, os, re, subprocess, sys, html as htmlmod

EDGE_CANDIDATES = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
]

# ---------- parsing ----------

def parse_front_matter(text):
    m = re.match(r"^---\n(.*?)\n---\n", text, re.S)
    if not m:
        return {}, text
    meta = {}
    for line in m.group(1).splitlines():
        if ":" in line and not line.startswith(" ") and not line.startswith("-"):
            k, v = line.split(":", 1)
            meta[k.strip()] = v.strip().strip('"')
    return meta, text[m.end():]

def parse_blocks(body):
    """Yield (kind, payload) blocks from the constrained markdown."""
    lines = body.split("\n")
    i, blocks = 0, []
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue
        if re.fullmatch(r"-{3,}", line.strip()):
            blocks.append(("hr", None)); i += 1; continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip())); i += 1; continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip())); i += 1; continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip())); i += 1; continue
        if re.fullmatch(r"_{8,}", line.strip()):
            blocks.append(("fill", None)); i += 1; continue
        if line.startswith("> "):
            quote = []
            while i < len(lines) and lines[i].startswith(">"):
                quote.append(lines[i][1:].lstrip()); i += 1
            blocks.append(("callout", " ".join(q for q in quote if q)))
            continue
        if line.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows)); continue
        if re.match(r"- \[ \] ", line):
            items = []
            while i < len(lines) and re.match(r"- \[ \] ", lines[i]):
                items.append(lines[i][6:].strip()); i += 1
            blocks.append(("tasks", items)); continue
        if line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip()); i += 1
            blocks.append(("bullets", items)); continue
        # numbered items are hand-numbered in the source; keep them as paragraphs
        para = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
                r"(#|\||> |- |-{3,}$|_{8,}$|\d+\. \*\*)", lines[i].strip()):
            para.append(lines[i].strip()); i += 1
        blocks.append(("p", " ".join(para)))
    return blocks

# ---------- inline formatting ----------

def inline_html(text):
    out = htmlmod.escape(text, quote=False)
    out = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", out)
    out = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<em>\1</em>", out)
    out = re.sub(r"`([^`]+)`", r"<code>\1</code>", out)
    return out

def inline_runs(text):
    """Split into (text, bold, italic) runs for docx."""
    runs, pattern = [], re.compile(r"(\*\*.+?\*\*|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        tok = m.group(0)
        if tok.startswith("**"):
            runs.append((tok[2:-2], True, False))
        else:
            runs.append((tok[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return [(t.replace("`", ""), b, it) for t, b, it in runs if t]

# ---------- HTML ----------

CSS = """
:root { --night:#0D1B2A; --ink:#14273A; --gold:#F0A500; --muted:#5A6B7A; --line:#D8DEE4; --boxbg:#F7F5F0; }
* { box-sizing: border-box; }
body { font-family: 'DM Sans', 'Segoe UI', Arial, sans-serif; color: var(--ink);
       margin: 0 auto; max-width: 760px; padding: 24px 20px; line-height: 1.65; font-size: 15.5px; background: #fff; }
h1, h2, h3 { font-family: 'Syne', 'Segoe UI', Arial, sans-serif; color: var(--night); line-height: 1.2; }
h1 { font-size: 1.9em; border-bottom: 3px solid var(--gold); padding-bottom: 8px; margin-top: 1.6em; }
h2 { font-size: 1.45em; margin-top: 1.5em; }
h3 { font-size: 1.12em; margin-top: 1.3em; }
p { margin: 0.7em 0; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 0.92em; }
th { background: var(--ink); color: #F2EFE9; text-align: left; }
th, td { border: 1px solid var(--line); padding: 7px 9px; vertical-align: top; }
td:empty::after { content: "\\00a0"; }
tr { break-inside: avoid; }
ul { margin: 0.6em 0; padding-left: 1.4em; }
li { margin: 0.3em 0; }
.tasks { list-style: none; padding-left: 0.2em; }
.tasks li::before { content: "\\2610\\00a0\\00a0"; color: var(--muted); }
.callout { background: var(--boxbg); border-left: 4px solid var(--gold);
           padding: 12px 16px; margin: 1.1em 0; border-radius: 0 6px 6px 0; break-inside: avoid; }
.fill { border-bottom: 1.2px solid #9FB0BE; height: 2.2em; margin: 0.4em 0 1em; }
hr { border: none; border-top: 1px solid var(--line); margin: 2em 0; }
code { background: #F0F2F4; padding: 1px 5px; border-radius: 4px; font-size: 0.9em; }
.cover { text-align: left; padding: 40px 0 30px; }
.cover .wordmark { font-family: 'Syne', Arial, sans-serif; font-weight: 800; font-size: 1.6em; color: var(--night); }
.cover .wordmark span { color: var(--gold); }
.cover h1.title { font-size: 2.3em; border: none; margin: 40px 0 8px; }
.cover .subtitle { color: var(--muted); font-size: 1.15em; max-width: 34em; }
.cover .rule { width: 88px; border-top: 4px solid var(--gold); margin: 26px 0; }
.cover table { max-width: 560px; }
.toc { columns: 2; column-gap: 36px; font-size: 0.95em; }
.toc a { color: var(--ink); text-decoration: none; }
.toc .l1 { font-weight: 700; margin-top: 0.6em; }
@media print {
  body { max-width: none; padding: 0; font-size: 10.6pt; }
  @page { size: A4; margin: 17mm 16mm; }
  h1 { break-before: page; }
  h2.chapter { break-before: page; }
  .cover, .tocpage { break-after: page; }
  .cover h1.title { break-before: avoid; }
  .callout, table { break-inside: avoid; }
  a { color: inherit; text-decoration: none; }
}
"""

def build_html(meta, blocks):
    parts = []
    add = parts.append
    add("<!doctype html><html lang='en'><head><meta charset='utf-8'>")
    add("<meta name='viewport' content='width=device-width, initial-scale=1'>")
    add(f"<title>{htmlmod.escape(meta.get('title','')) }</title>")
    add("<link rel='preconnect' href='https://fonts.googleapis.com'>")
    add("<link href='https://fonts.googleapis.com/css2?family=Syne:wght@700;800&family=DM+Sans:opsz,wght@9..40,400;9..40,700&display=swap' rel='stylesheet'>")
    add(f"<style>{CSS}</style></head><body>")

    # Cover (light background — knowledge/branding/visual-language.md print rule)
    add("<div class='cover'>")
    add("<div class='wordmark'>en<span>aasan</span></div>")
    add("<div class='rule'></div>")
    add(f"<h1 class='title'>{htmlmod.escape(meta.get('title',''))}</h1>")
    add(f"<p class='subtitle'>{htmlmod.escape(meta.get('subtitle',''))}</p>")
    add("<table>")
    for label, key in [("Document ID","document_id"),("Version","version"),("Status","status"),
                       ("Author","author"),("Audience","audience"),
                       ("Estimated completion","estimated_completion_time"),
                       ("Last reviewed","last_reviewed"),("Review due","review_due")]:
        if meta.get(key):
            add(f"<tr><td><strong>{label}</strong></td><td>{htmlmod.escape(meta[key])}</td></tr>")
    add("</table>")
    add(f"<p class='subtitle'>{htmlmod.escape(meta.get('license',''))}</p>")
    add("</div>")

    # TOC from h1/h2
    add("<div class='tocpage'><h2>Contents</h2><div class='toc'>")
    n = 0
    for kind, payload in blocks:
        if kind in ("h1", "h2"):
            n += 1
            cls = "l1" if kind == "h1" else "l2"
            add(f"<div class='{cls}'><a href='#s{n}'>{inline_html(payload)}</a></div>")
    add("</div></div>")

    n = 0
    skipped_title = False
    for kind, payload in blocks:
        if kind == "h1":
            n += 1
            if not skipped_title:  # source's own H1 title duplicates the cover
                skipped_title = True
                add(f"<span id='s{n}'></span>")
                continue
            add(f"<h1 id='s{n}'>{inline_html(payload)}</h1>")
        elif kind == "h2":
            n += 1
            cls = " class='chapter'" if payload.lower().startswith("chapter") else " class='chapter'"
            add(f"<h2{cls} id='s{n}'>{inline_html(payload)}</h2>")
        elif kind == "h3":
            add(f"<h3>{inline_html(payload)}</h3>")
        elif kind == "p":
            add(f"<p>{inline_html(payload)}</p>")
        elif kind == "callout":
            add(f"<div class='callout'>{inline_html(payload)}</div>")
        elif kind == "bullets":
            add("<ul>" + "".join(f"<li>{inline_html(x)}</li>" for x in payload) + "</ul>")
        elif kind == "tasks":
            add("<ul class='tasks'>" + "".join(f"<li>{inline_html(x)}</li>" for x in payload) + "</ul>")
        elif kind == "table":
            head, *rows = payload
            add("<table><tr>" + "".join(f"<th>{inline_html(c)}</th>" for c in head) + "</tr>")
            for r in rows:
                r = r + [""] * (len(head) - len(r))
                add("<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in r) + "</tr>")
            add("</table>")
        elif kind == "fill":
            add("<div class='fill'></div>")
        elif kind == "hr":
            add("<hr>")
    add("</body></html>")
    return "\n".join(parts)

# ---------- PDF ----------

def build_pdf(html_path, pdf_path):
    edge = next((p for p in EDGE_CANDIDATES if os.path.exists(p)), None)
    if not edge:
        print("!! Edge not found — skipping PDF"); return False
    subprocess.run([edge, "--headless", "--disable-gpu", "--no-pdf-header-footer",
                    f"--print-to-pdf={os.path.abspath(pdf_path)}",
                    "file:///" + os.path.abspath(html_path).replace("\\", "/")],
                   capture_output=True, timeout=180)
    return os.path.exists(pdf_path)

# ---------- DOCX ----------

def build_docx(meta, blocks, path):
    import docx
    from docx.shared import Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY, GOLD, MUTED = RGBColor(0x14,0x27,0x3A), RGBColor(0xF0,0xA5,0x00), RGBColor(0x5A,0x6B,0x7A)
    doc = docx.Document()
    doc.core_properties.title = meta.get("title","")
    doc.core_properties.author = meta.get("author","")
    doc.core_properties.comments = meta.get("document_id","")

    st = doc.styles["Normal"];  st.font.name = "Calibri"; st.font.size = Pt(10.5)
    for name, size in [("Title",24),("Heading 1",17),("Heading 2",13.5),("Heading 3",11.5)]:
        s = doc.styles[name]; s.font.name = "Calibri"; s.font.bold = True
        s.font.size = Pt(size); s.font.color.rgb = NAVY

    def shade(cell, hexfill):
        el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hexfill)
        cell._tc.get_or_add_tcPr().append(el)

    def bottom_border(par):
        pPr = par._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8"); b.set(qn("w:color"), "9FB0BE")
        pBdr.append(b); pPr.append(pBdr)

    def add_runs(par, text):
        for t, b, i in inline_runs(text):
            r = par.add_run(t); r.bold, r.italic = b, i

    def page_break():
        doc.add_page_break()

    # Cover — light background per the print-friendliness rule
    p = doc.add_paragraph(); r1 = p.add_run("en"); r1.bold = True; r1.font.size = Pt(20); r1.font.color.rgb = NAVY
    r2 = p.add_run("aasan"); r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = GOLD
    t = doc.add_paragraph(style="Title"); t.add_run(meta.get("title",""))
    sub = doc.add_paragraph(); sr = sub.add_run(meta.get("subtitle","")); sr.italic = True; sr.font.color.rgb = MUTED
    doc.add_paragraph()
    metatbl = [("Document ID", meta.get("document_id","")),("Version", meta.get("version","")),
               ("Status", meta.get("status","")),("Author", meta.get("author","")),
               ("Estimated completion", meta.get("estimated_completion_time","")),
               ("Last reviewed", meta.get("last_reviewed","")),("Review due", meta.get("review_due",""))]
    tab = doc.add_table(rows=0, cols=2); tab.style = "Table Grid"
    for k, v in metatbl:
        if not v: continue
        row = tab.add_row(); row.cells[0].text = k; row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Twips(2400); row.cells[1].width = Twips(6600)
    doc.add_paragraph()
    lic = doc.add_paragraph(); lr = lic.add_run(meta.get("license","")); lr.font.size = Pt(8.5); lr.font.color.rgb = MUTED
    page_break()

    skipped_title = False
    for kind, payload in blocks:
        if kind == "h1":
            if not skipped_title: skipped_title = True; continue
            page_break(); h = doc.add_paragraph(style="Heading 1"); add_runs(h, payload)
        elif kind == "h2":
            page_break(); h = doc.add_paragraph(style="Heading 2"); add_runs(h, payload)
        elif kind == "h3":
            h = doc.add_paragraph(style="Heading 3"); add_runs(h, payload)
        elif kind == "p":
            add_runs(doc.add_paragraph(), payload)
        elif kind == "callout":
            box = doc.add_table(rows=1, cols=1); box.style = "Table Grid"
            cell = box.rows[0].cells[0]; cell.width = Twips(9000); shade(cell, "F7F5F0")
            add_runs(cell.paragraphs[0], payload)
            doc.add_paragraph()
        elif kind == "bullets":
            for it in payload:
                add_runs(doc.add_paragraph(style="List Bullet"), it)
        elif kind == "tasks":
            for it in payload:
                par = doc.add_paragraph(); par.add_run("☐  ")
                add_runs(par, it)
        elif kind == "table":
            head, *rows = payload
            ncols = len(head)
            tbl = doc.add_table(rows=1, cols=ncols); tbl.style = "Table Grid"
            width = Twips(int(9000 / ncols))
            for j, c in enumerate(head):
                cell = tbl.rows[0].cells[j]; cell.width = width; shade(cell, "14273A")
                par = cell.paragraphs[0]
                for tt, b, i in inline_runs(c):
                    r = par.add_run(tt); r.bold = True; r.font.color.rgb = RGBColor(0xF2,0xEF,0xE9)
            for rdata in rows:
                rdata = rdata + [""] * (ncols - len(rdata))
                row = tbl.add_row()
                for j, c in enumerate(rdata[:ncols]):
                    row.cells[j].width = width
                    if c: add_runs(row.cells[j].paragraphs[0], c)
            doc.add_paragraph()
        elif kind == "fill":
            par = doc.add_paragraph(); par.add_run(" "); bottom_border(par)
            doc.add_paragraph()
        elif kind == "hr":
            pass
    doc.save(path)
    return True

# ---------- main ----------

def main():
    src = sys.argv[1]
    skip_pdf = "--skip-pdf" in sys.argv
    text = io.open(src, encoding="utf-8").read()
    meta, body = parse_front_matter(text)
    blocks = parse_blocks(body)
    doc_id = meta.get("document_id", "untitled")

    root = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "output"))
    html_path = os.path.join(root, "html", f"{doc_id}.html")
    pdf_path  = os.path.join(root, "pdf",  f"{doc_id}.pdf")
    docx_path = os.path.join(root, "docx", f"{doc_id}.docx")

    io.open(html_path, "w", encoding="utf-8", newline="\n").write(build_html(meta, blocks))
    print("html :", html_path)
    if not skip_pdf:
        print("pdf  :", pdf_path if build_pdf(html_path, pdf_path) else "FAILED")
    print("docx :", docx_path if build_docx(meta, blocks, docx_path) else "FAILED")
    print(f"blocks: {len(blocks)}")

if __name__ == "__main__":
    main()
